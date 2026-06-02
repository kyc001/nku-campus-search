from __future__ import annotations

from pathlib import Path
from urllib.parse import urlparse

from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas

from .config import ATTACHMENT_DIR, CUSTOM_DICT_PATH, DOCUMENTS_PATH, SEEDS_PATH, ensure_dirs
from .crawler import infer_tags
from .indexer import build_anchor_texts, build_index
from .models import Document, utc_now_iso
from .parser import parse_attachment
from .snapshot import save_html_snapshot, url_id
from .storage import save_documents


BASE_PAGES = [
    (
        "https://news.nankai.edu.cn/ywsd/system/2026/05/14/030066001.shtml",
        "南开大学举行人工智能与信息检索交叉论坛",
        "南开大学计算机学院举办人工智能与信息检索交叉论坛，来自八里台和津南校区的师生围绕搜索引擎、向量空间模型、链接分析和大模型应用展开讨论。",
    ),
    (
        "https://news.nankai.edu.cn/ywsd/system/2026/05/12/030066002.shtml",
        "南开新闻网报道伯苓班拔尖创新人才培养",
        "伯苓班持续推进基础学科拔尖创新人才培养，学校发布了课程建设、导师制和科研训练的新进展。",
    ),
    (
        "https://www.nankai.edu.cn/xxgk/list.htm",
        "南开大学学校概况",
        "南开大学是教育部直属重点综合性大学，秉承允公允能、日新月异校训，设有八里台校区、津南校区和泰达学院。",
    ),
    (
        "https://cc.nankai.edu.cn/2026/0518/c13619a550001/page.htm",
        "计算机学院关于操作系统课程考试安排的通知",
        "计算机学院发布操作系统、计算机网络和数据库系统课程考试安排。请本科生登录教务系统查看考场和时间。",
    ),
    (
        "https://cc.nankai.edu.cn/teachers/wenliang.htm",
        "温良教授主页",
        "温良，南开大学计算机学院教授，研究方向包括信息检索、Web 搜索、图学习和自然语言处理。欢迎对计网和搜索技术感兴趣的同学联系。",
    ),
    (
        "https://jwc.nankai.edu.cn/2026/0516/courses.htm",
        "教务部关于 2026 年夏季学期选课的通知",
        "教务部通知本科生和研究生完成夏季学期选课，涉及专业必修课、通识课、体育课和重修报名。",
    ),
    (
        "https://lib.nankai.edu.cn/2026/database.htm",
        "图书馆新增中外文数据库资源",
        "南开大学图书馆新增若干学术数据库和电子期刊资源，支持校园网访问和校外统一身份认证访问。",
    ),
    (
        "https://xg.nankai.edu.cn/2026/jobfair.htm",
        "学生就业指导中心春季招聘宣讲会安排",
        "学生就业指导中心发布春季招聘宣讲会安排，覆盖软件开发、金融科技、教育培训等方向。",
    ),
    (
        "https://math.nankai.edu.cn/2026/seminar.htm",
        "数学科学学院偏微分方程学术讲座",
        "数学科学学院邀请校内外专家举办偏微分方程和科学计算学术讲座，欢迎师生参加。",
    ),
    (
        "https://12club.nankai.edu.cn/2026/comic.htm",
        "12 Club 动漫社团活动月",
        "南开 12 Club 动漫社团举办活动月，包含漫画分享、动画放映、社团招新和文创展示。",
    ),
]

DEFAULT_SEEDS = [
    "https://www.nankai.edu.cn/",
    "https://news.nankai.edu.cn/",
    "https://cc.nankai.edu.cn/",
    "https://jwc.nankai.edu.cn/",
    "https://lib.nankai.edu.cn/",
    "https://xg.nankai.edu.cn/",
    "https://math.nankai.edu.cn/",
    "https://physics.nankai.edu.cn/",
    "https://chem.nankai.edu.cn/",
    "https://sky.nankai.edu.cn/",
    "https://cs.nankai.edu.cn/",
    "https://ai.nankai.edu.cn/",
    "https://software.nankai.edu.cn/",
    "https://stat.nankai.edu.cn/",
    "https://econ.nankai.edu.cn/",
    "https://bs.nankai.edu.cn/",
    "https://law.nankai.edu.cn/",
    "https://history.nankai.edu.cn/",
    "https://phil.nankai.edu.cn/",
    "https://foreign.nankai.edu.cn/",
    "https://environment.nankai.edu.cn/",
    "https://medicine.nankai.edu.cn/",
    "https://pharmacy.nankai.edu.cn/",
    "https://mse.nankai.edu.cn/",
    "https://ceo.nankai.edu.cn/",
    "https://graduate.nankai.edu.cn/",
    "https://yzb.nankai.edu.cn/",
    "https://zsb.nankai.edu.cn/",
    "https://job.nankai.edu.cn/",
    "https://oldnews.nankai.edu.cn/",
    "https://12club.nankai.edu.cn/",
]


def _expanded_pages(total: int = 2000) -> list[tuple[str, str, str]]:
    pages = list(BASE_PAGES)
    topics = [
        ("news.nankai.edu.cn", "南开新闻网", "新闻", "学校发布南开新闻与综合报道，内容涉及人才培养、科研进展和校园文化。"),
        ("cc.nankai.edu.cn", "计算机学院", "院系", "计算机学院发布课程通知、教师信息、信息检索、数据库系统和计算机网络相关资源。"),
        ("jwc.nankai.edu.cn", "教务部", "教务", "教务部发布选课、考试、培养方案、课程建设和教学运行通知。"),
        ("lib.nankai.edu.cn", "图书馆", "图书馆", "图书馆发布数据库、电子期刊、学术资源和统一身份认证访问说明。"),
        ("xg.nankai.edu.cn", "学生就业指导中心", "招聘", "学生就业指导中心发布招聘宣讲、实习岗位、就业服务和学生发展通知。"),
        ("math.nankai.edu.cn", "数学科学学院", "学术", "数学科学学院发布学术讲座、科研训练、偏微分方程和科学计算信息。"),
        ("12club.nankai.edu.cn", "12 Club", "文体", "南开 12 Club 发布动漫社团活动、文创展示和社团招新通知。"),
    ]
    names = ["温良", "王明", "李华", "张宁", "赵一", "陈津", "刘开", "周南"]
    idx = 0
    while len(pages) < total:
        host, site_name, tag, body_seed = topics[idx % len(topics)]
        name = names[idx % len(names)]
        year = 2026
        month = (idx % 12) + 1
        day = (idx % 27) + 1
        if host == "cc.nankai.edu.cn" and idx % 11 == 0:
            title = f"{name}教师主页与信息检索研究方向"
            body = f"{name}老师来自南开大学{site_name}，研究方向包括信息检索、Web 搜索、向量空间模型、PageRank 和计网。"
            path = f"/teachers/demo-{idx:04d}.htm"
        else:
            title = f"{site_name}{tag}资源第 {idx + 1:04d} 条"
            body = f"{body_seed} 本页面属于南开大学校内资源，支持站内查询、短语查询、通配查询、网页快照和个性化推荐。"
            path = f"/2026/{month:02d}{day:02d}/demo-{idx:04d}.htm"
        pages.append((f"https://{host}{path}", title, body))
        idx += 1
    return pages


def _html_doc(title: str, body: str, links: list[str]) -> str:
    link_html = "\n".join(f'<li><a href="{link}">{Path(urlparse(link).path).stem or link}</a></li>' for link in links)
    return f"""<!doctype html>
<html lang="zh-CN">
<head><meta charset="utf-8"><title>{title}</title></head>
<body>
<header><h1>{title}</h1></header>
<main><p>{body}</p><section><h2>相关链接</h2><ul>{link_html}</ul></section></main>
</body>
</html>"""


def _make_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=A4)
    try:
        font_path = Path("C:/Windows/Fonts/msyh.ttc")
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("MSYH", str(font_path)))
            c.setFont("MSYH", 14)
    except Exception:
        c.setFont("Helvetica", 12)
    lines = [
        "南开大学信息检索课程实验说明",
        "本 PDF 附件用于演示文档查询 filetype:pdf。",
        "内容包含向量空间模型、倒排索引、PageRank 链接分析、网页快照和查询日志。",
        "短语查询示例：南开 大学；通配查询示例：计?、温*。",
    ]
    y = 780
    for line in lines:
        c.drawString(72, y, line)
        y -= 30
    c.save()


def _make_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("计算机学院研究生课程培养方案", level=1)
    doc.add_paragraph("本 DOCX 附件用于演示文档查询 filetype:docx。")
    doc.add_paragraph("培养方案包含信息检索、机器学习、数据库系统、计算机网络等课程。")
    doc.add_paragraph("教师和研究生用户在个性化排序中会更容易看到学术与文档资源。")
    doc.save(str(path))


def _make_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "招聘宣讲"
    ws.append(["日期", "单位", "地点", "关键词"])
    ws.append(["2026-05-21", "南开软件联合实验室", "津南校区", "招聘 实习 软件开发"])
    ws.append(["2026-05-22", "信息检索研究组", "八里台校区", "学术 搜索 引擎"])
    wb.save(path)


def create_demo_dataset(total_pages: int = 2000) -> list[Document]:
    ensure_dirs()
    CUSTOM_DICT_PATH.write_text(
        "\n".join(["南开", "南开大学", "八里台", "津南", "伯苓班", "信息检索", "向量空间模型", "PageRank", "计网", "计算机网络"]),
        encoding="utf-8",
    )
    SEEDS_PATH.write_text("\n".join(DEFAULT_SEEDS) + "\n", encoding="utf-8")

    pages = _expanded_pages(total_pages)
    urls = [page[0] for page in pages]
    documents: list[Document] = []
    for idx, (url, title, body) in enumerate(pages):
        links = [urls[(idx + 1) % len(urls)], urls[(idx + 3) % len(urls)]]
        if "cc.nankai.edu.cn" in url:
            links.append("https://cc.nankai.edu.cn/files/course-plan.docx")
        if "jwc.nankai.edu.cn" in url:
            links.append("https://jwc.nankai.edu.cn/files/ir-homework.pdf")
        if "xg.nankai.edu.cn" in url:
            links.append("https://xg.nankai.edu.cn/files/jobfair.xlsx")
        html = _html_doc(title, body, links)
        snapshot = save_html_snapshot(url, html)
        documents.append(
            Document(
                id=url_id(url),
                url=url,
                title=title,
                body=body,
                host=urlparse(url).netloc,
                doc_type="html",
                outlinks=links,
                crawled_at=utc_now_iso(),
                snapshot_path=snapshot,
                tags=infer_tags(title + body),
            )
        )

    attachments = [
        ("https://jwc.nankai.edu.cn/files/ir-homework.pdf", "pdf", ATTACHMENT_DIR / "ir-homework.pdf", _make_pdf),
        ("https://cc.nankai.edu.cn/files/course-plan.docx", "docx", ATTACHMENT_DIR / "course-plan.docx", _make_docx),
        ("https://xg.nankai.edu.cn/files/jobfair.xlsx", "xlsx", ATTACHMENT_DIR / "jobfair.xlsx", _make_xlsx),
    ]
    for url, doc_type, path, maker in attachments:
        maker(path)
        body = parse_attachment(path, doc_type)
        documents.append(
            Document(
                id=url_id(url),
                url=url,
                title=path.stem,
                body=body,
                host=urlparse(url).netloc,
                doc_type=doc_type,
                attachment_url=url,
                attachment_path=str(path),
                file_size=path.stat().st_size,
                crawled_at=utc_now_iso(),
                tags=infer_tags(body),
            )
        )

    documents = build_anchor_texts(documents)
    save_documents(documents, DOCUMENTS_PATH)
    build_index(documents)
    return documents
