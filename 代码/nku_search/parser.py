from __future__ import annotations

import re
from pathlib import Path
from urllib.parse import urljoin, urlparse

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from openpyxl import load_workbook
from pypdf import PdfReader

from .text import normalize_text

DOC_SUFFIXES = {".pdf", ".doc", ".docx", ".xls", ".xlsx"}


def clean_url(url: str) -> str:
    parsed = urlparse(url)
    scheme = parsed.scheme.lower() or "https"
    host = parsed.netloc.lower()
    path = parsed.path or "/"
    return f"{scheme}://{host}{path}" + (f"?{parsed.query}" if parsed.query else "")


def guess_doc_type(url: str, content_type: str = "") -> str:
    suffix = Path(urlparse(url).path).suffix.lower().lstrip(".")
    if suffix in {"pdf", "doc", "docx", "xls", "xlsx"}:
        return suffix
    content_type = content_type.lower()
    if "pdf" in content_type:
        return "pdf"
    if "word" in content_type:
        return "docx"
    if "excel" in content_type or "spreadsheet" in content_type:
        return "xlsx"
    return "html"


def parse_html(html: str, base_url: str) -> tuple[str, str, list[tuple[str, str]]]:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript", "svg"]):
        tag.decompose()
    title = normalize_text(soup.title.get_text(" ", strip=True) if soup.title else "")
    main = soup.find("main") or soup.find("article") or soup.body or soup
    body = normalize_text(main.get_text(" ", strip=True))
    links: list[tuple[str, str]] = []
    for a in soup.find_all("a", href=True):
        href = a.get("href", "").strip()
        if not href or href.startswith(("javascript:", "mailto:", "tel:")):
            continue
        try:
            absolute = clean_url(urljoin(base_url, href))
        except ValueError:
            continue
        anchor = normalize_text(a.get_text(" ", strip=True))
        links.append((absolute, anchor))
    return title, body, links


def parse_pdf(path: Path, max_pages: int = 30) -> str:
    reader = PdfReader(str(path))
    pages = reader.pages[:max_pages]
    return normalize_text("\n".join(page.extract_text() or "" for page in pages))


def parse_docx(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" ".join(cell.text for cell in row.cells))
    return normalize_text("\n".join(parts))


def parse_xlsx(path: Path, max_rows: int = 300) -> str:
    wb = load_workbook(path, read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(ws.title)
        for idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
            if idx > max_rows:
                break
            values = [str(value) for value in row if value is not None]
            if values:
                parts.append(" ".join(values))
    return normalize_text("\n".join(parts))


def parse_attachment(path: Path, doc_type: str) -> str:
    doc_type = doc_type.lower()
    try:
        if doc_type == "pdf":
            return parse_pdf(path)
        if doc_type == "docx":
            return parse_docx(path)
        if doc_type == "xlsx":
            return parse_xlsx(path)
    except Exception as exc:
        return f"附件解析失败: {type(exc).__name__}: {exc}"
    return "旧版 doc/xls 文件已保存，建议安装 Apache Tika 或 LibreOffice 后扩展解析。"


def looks_like_attachment(url: str) -> bool:
    return Path(urlparse(url).path).suffix.lower() in DOC_SUFFIXES


def safe_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.-]+", "_", name).strip("_")[:120] or "attachment"
